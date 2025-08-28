import uuid
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.pdfgen import canvas
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from io import BytesIO
from utils.mongo_utils import find_student_by_pin, update_student_hallticket
import os

# The _replace_placeholders_in_doc and _extract_subject_rows functions remain the same.
def _replace_placeholders_in_doc(doc: Document, replacements: dict):
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)


def _extract_subject_rows(doc: Document):
    rows = []
    for table in doc.tables:
        cells_text = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        header_found = any("Date" in c or "Subject" in c for r in cells_text for c in r)
        if header_found:
            for r in cells_text:
                if any("Date" in c or "Subject" in c for c in r):
                    continue
                if len(r) >= 2 and r[0] and r[1]:
                    rows.append((r[0], r[1]))
            if rows:
                return rows
    return rows


def generate_hallticket_pdf(pin: str):
    student = find_student_by_pin(pin)
    if not student:
        return None

    hallticket_id = student.get("hallticket_id")
    if not hallticket_id:
        hallticket_id = str(uuid.uuid4()).split("-")[0].upper()
        update_student_hallticket(pin, None, hallticket_id=hallticket_id)

    dept = student.get("department", "").upper()
    year = int(student.get("year", 0))

    template_map = {
        (3, "AIML"): "./docs/template_3_AIML.docx",
        (3, "CS"): "./docs/template_3_CS.docx",
        (4, "AIML"): "./docs/template_4_AIML.docx",
        (4, "CS"): "./docs/template_4_CS.docx",
        (4, "DS"): "./docs/template_4_DS.docx",
        (3, "DS"): "./docs/template_3_DS.docx",
        (3, "ECE"): "./docs/template_3_ECE.docx",
        (4, "ECE"): "./docs/template_4_ECE.docx",
        (2, "AIML"): "./docs/template_2_AIML.docx",
        (2, "CS"): "./docs/template_2_CS.docx",
        (2, "DS"): "./docs/template_2_DS.docx",
        (2, "ECE"): "./docs/template_2_ECE.docx"
    }
    template = template_map.get((year, dept))
    if not template or not os.path.isfile(template):
        return None

    doc = Document(template)
    replacements = {
        "{{NAME}}": student["name"],
        "{{PIN}}": student["pin"],
        "{{ID}}": hallticket_id
    }
    _replace_placeholders_in_doc(doc, replacements)

    subjects = _extract_subject_rows(doc) or [
        ("18-08-2025", "Subject A"),
        ("19-08-2025", "Subject B"),
        ("20-08-2025", "Subject C"),
    ]

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    if year == 2:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  
        logo_path = os.path.join(base_dir, "assets", "ggu-logo.png")
    else:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  
        logo_path = os.path.join(base_dir, "assets", "giet-logo.jpg")
    # ==================== MODIFICATION START ====================

    # 1. ADD IMAGE TO THE TOP OF THE PDF
    # IMPORTANT: Replace "path/to/your/logo.png" with the actual path to your image file.
    # Make sure the image file is accessible where you run the script.
    try:
          # <-- using your assets folder
        logo_width = 40 * mm
        logo_height = 40 * mm
        c.drawImage(
            logo_path,
            (width - logo_width) / 2,  # center horizontally
            height - 55 * mm,          # adjust Y position a little higher
            width=logo_width,
            height=logo_height,
            preserveAspectRatio=True,
            mask="auto"
        )
    except Exception as e:
        print(f"Could not draw image. Error: {e}")
        # If the image fails to load, the PDF will be generated without it.

    # 2. ADJUST STARTING Y-COORDINATE FOR THE TEXT
    # This moves all subsequent content down to make space for the logo.
    y = height - 60 * mm  # Original value was height - 35 * mm

    # ===================== MODIFICATION END =====================

    # Styles
    title_style = ParagraphStyle(name="Title", fontName="Times-Bold", fontSize=16, alignment=TA_CENTER, leading=15)
    subtitle_style = ParagraphStyle(name="Sub", fontName="Times-Roman", fontSize=12, alignment=TA_CENTER, leading=13)

    if year == 2:
        para = Paragraph("GODAVARI GLOBAL UNIVERSITY", title_style)
        w, h = para.wrap(width - 80, 100)
        para.drawOn(c, 40, y)
        y -= h + 4

    else:
        para = Paragraph("GODAVARI INSTITUTE OF ENGINEERING & TECHNOLOGY", title_style)
        w, h = para.wrap(width - 80, 100)
        para.drawOn(c, 40, y)
        y -= h + 4

        para = Paragraph("Approved By AICTE | NAAC ‘A++’ | Recognized by UGC,", subtitle_style)
        w, h = para.wrap(width - 80, 100)
        para.drawOn(c, 40, y)
        y -= h + 2

        para = Paragraph("U/Sec. 2(f) & 12(B) | Permanently Affiliated to JNTUK, Kakinada", subtitle_style)
        w, h = para.wrap(width - 80, 100)
        para.drawOn(c, 40, y)
        y -= h

    c.setStrokeColor(colors.black)
    c.setLineWidth(0.8)
    c.line(40, y, width - 40, y)
    y -= 12

    if dept == "ECE":
        dept_line = "Electronics & Communication Engineering"
    else:
        dept_line = f"Computer Science & Engineering({student['department']})"
    c.setFont("Times-Bold", 12)
    c.drawCentredString(width / 2, y, dept_line)
    y -= 16

    if year == 2:
        exam_title = "Hall Ticket: Mid-1 Examinations, II B.Tech I Sem (A.Y: 2025-26)"
        exam_time = "Time: 2:45 pm - 4:15pm"
    elif year == 3:
        exam_title = "Hall Ticket: Mid-1 Examinations, III B.Tech I Sem (A.Y: 2025-26)"
        exam_time = "Time: Objective Exam – 10:50 am – 11:00 am, Descriptive Exam: 11:00 am – 12:30 pm"
    else:
        exam_title = "Hall Ticket: Mid-1 Examinations, IV B.Tech I Sem (A.Y: 2025-26)"
        exam_time = "Time: Objective Exam – 2:30 pm – 2:40 pm, Descriptive Exam: 2:40 pm – 04:10 pm"
    c.setFont("Times-Bold", 12)
    c.drawCentredString(width / 2, y, exam_title)
    y -= 20

    student_info = [
        ["Hall Ticket No:", student["pin"], "ID:", hallticket_id],
        ["Name:", student["name"], "", ""],
        ["Branch:", student["department"], "Year:", str(student["year"])],
    ]

    table = Table(student_info, colWidths=[35*mm, 50*mm, 25*mm, 50*mm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("SPAN", (1, 1), (3, 1)),
        ("ALIGN", (0, 0), (0, -1), "LEFT"),
        ("ALIGN", (2, 0), (2, -1), "LEFT"),
        ("ALIGN", (1, 0), (1, 0), "LEFT"),
        ("ALIGN", (3, 0), (3, 0), "LEFT"),
        ("ALIGN", (1, 1), (3, 1), "LEFT"),
        ("ALIGN", (1, 2), (1, 2), "LEFT"),
        ("ALIGN", (3, 2), (3, 2), "LEFT"),
        ("FONTNAME", (0, 0), (0, -1), "Times-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Times-Bold"),
    ]))

    tw, th = table.wrapOn(c, width - 80, y)
    table.drawOn(c, 40, y - th)
    y -= th + 18

    c.setFont("Times-Roman", 11)
    c.drawString(40, y, exam_time)
    y -= 18

    c.setFont("Times-Bold", 12)
    c.drawString(40, y, "Exam Schedule:")
    y -= 16

    table_data = [["Date", "Subject Name", "Signature of Invigilator"]] + subjects
    col_widths = [40*mm, width - 40*2 - 90*mm, 50*mm]
    subj_table = Table(table_data, colWidths=col_widths, repeatRows=1)
    subj_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
    ]))
    tw, th = subj_table.wrapOn(c, width - 80, y)
    subj_table.drawOn(c, 40, y - th)
    y -= th + 65

    c.setFont("Times-Roman", 11)
    c.drawString(60, y, "Signature of Student")
    c.drawRightString(width - 60, y, "Signature of HOD/Class Teacher")
    c.showPage()
    c.save()

    buffer.seek(0)
    update_student_hallticket(pin, None, hallticket_id=hallticket_id)

    return buffer